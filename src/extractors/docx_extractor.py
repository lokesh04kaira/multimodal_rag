from pathlib import Path
from src.utils import clean_text

def extract_docx(path: str) -> str:
    p = Path(path)
    if not p.exists() or not p.is_file():
        return ""
    try:
        import docx
        d = docx.Document(str(p))
        parts = []
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text or "" for c in row.cells]
                parts.append(" | ".join(cells))
        for para in d.paragraphs:
            parts.append(para.text or "")
        return clean_text("\n".join(parts))
    except Exception:
        return ""
